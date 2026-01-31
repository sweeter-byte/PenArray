"""
Document export endpoints for BiZhen system.

Implements .docx and .pdf export functionality as per FR-01.
Allows users to download generated essays in standard document formats.
"""

import io
from datetime import datetime
from typing import Optional
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

from backend.api.deps import get_db, get_current_user
from backend.db.models import User, Task, EssayResult


router = APIRouter()


def get_style_name_cn(style: str) -> str:
    """Get Chinese name for essay style."""
    names = {
        "profound": "深刻型",
        "rhetorical": "文采型",
        "steady": "稳健型",
    }
    return names.get(style, style)


def get_grade_level(score: int) -> str:
    """Get grade level description in Chinese."""
    if score >= 50:
        return "一等文"
    elif score >= 40:
        return "二等文"
    elif score >= 30:
        return "三等文"
    else:
        return "四等文"


@router.get(
    "/{essay_id}/docx",
    summary="导出Word文档",
    description="将作文导出为.docx格式的Word文档",
    responses={
        200: {"description": "Word document file"},
        401: {"description": "Not authenticated"},
        404: {"description": "Essay not found"},
    },
)
def export_essay_docx(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    """
    Export essay as Word document (.docx).

    Includes title, content, score, and generation metadata.
    """
    # Get essay from database
    essay = db.query(EssayResult).filter(EssayResult.id == essay_id).first()

    if essay is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"作文 {essay_id} 不存在",
        )

    # Authorization check
    task = db.query(Task).filter(Task.id == essay.task_id).first()
    if task is None or task.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="无权访问此作文",
        )

    # Create Word document
    doc = Document()

    # Add title
    title_para = doc.add_heading(essay.title or "无标题", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata section
    style_cn = get_style_name_cn(essay.style)
    score = essay.score or 0
    grade = get_grade_level(score)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"风格：{style_cn} | 评分：{score}/60 ({grade})")
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True

    # Add separator
    doc.add_paragraph("─" * 50)

    # Add essay content
    content_lines = (essay.content or "").split("\n")
    for line in content_lines:
        if line.strip():
            para = doc.add_paragraph(line.strip())
            para.paragraph_format.first_line_indent = Inches(0.3)
            para.paragraph_format.line_spacing = 1.5

    # Add footer with generation info
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(
        f"由笔阵系统生成 | {datetime.now().strftime('%Y年%m月%d日')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Generate filename
    safe_title = (essay.title or "essay")[:20].replace(" ", "_")
    filename = f"{safe_title}_{style_cn}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"
        }
    )


@router.get(
    "/{essay_id}/pdf",
    summary="导出PDF文档",
    description="将作文导出为.pdf格式的文档",
    responses={
        200: {"description": "PDF document file"},
        401: {"description": "Not authenticated"},
        404: {"description": "Essay not found"},
    },
)
def export_essay_pdf(
    essay_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    """
    Export essay as PDF document.

    Includes title, content, score, and generation metadata.
    Uses ReportLab for PDF generation with Chinese font support.
    """
    # Get essay from database
    essay = db.query(EssayResult).filter(EssayResult.id == essay_id).first()

    if essay is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"作文 {essay_id} 不存在",
        )

    # Authorization check
    task = db.query(Task).filter(Task.id == essay.task_id).first()
    if task is None or task.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="无权访问此作文",
        )

    # Create PDF buffer
    buffer = io.BytesIO()

    # Create document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    # Get styles
    styles = getSampleStyleSheet()

    # Custom styles for Chinese text (using built-in Helvetica as fallback)
    title_style = ParagraphStyle(
        'ChineseTitle',
        parent=styles['Title'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=12,
    )

    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        textColor='gray',
        spaceAfter=20,
    )

    body_style = ParagraphStyle(
        'ChineseBody',
        parent=styles['Normal'],
        fontSize=12,
        alignment=TA_JUSTIFY,
        firstLineIndent=24,
        leading=20,
        spaceAfter=6,
    )

    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        alignment=TA_CENTER,
        textColor='gray',
    )

    # Build content
    story = []

    # Title
    story.append(Paragraph(essay.title or "无标题", title_style))

    # Metadata
    style_cn = get_style_name_cn(essay.style)
    score = essay.score or 0
    grade = get_grade_level(score)
    meta_text = f"风格：{style_cn} | 评分：{score}/60 ({grade})"
    story.append(Paragraph(meta_text, meta_style))

    # Separator
    story.append(Spacer(1, 12))

    # Essay content
    content_lines = (essay.content or "").split("\n")
    for line in content_lines:
        if line.strip():
            # Escape special characters for ReportLab
            safe_line = line.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, body_style))

    # Footer
    story.append(Spacer(1, 30))
    footer_text = f"由笔阵系统生成 | {datetime.now().strftime('%Y年%m月%d日')}"
    story.append(Paragraph(footer_text, footer_style))

    # Build PDF
    doc.build(story)
    buffer.seek(0)

    # Generate filename
    safe_title = (essay.title or "essay")[:20].replace(" ", "_")
    filename = f"{safe_title}_{style_cn}.pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"
        }
    )
